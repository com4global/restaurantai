"""
Menu Extractor — Extract structured menu data from images using Gemini Vision.

Supports: JPEG, PNG, WebP images of restaurant menus.
Returns: {categories: [{name, items: [{name, description, price_cents}]}]}
"""

import base64
import json as _json
import os
import urllib.parse
import urllib.request
from dotenv import load_dotenv


EXTRACTION_PROMPT = """You are extracting a restaurant menu from an uploaded image (photo of a printed menu, screenshot, etc).
Return JSON:
{
  "restaurant_name": "...",
  "categories": [
    {
      "name": "Category Name",
      "items": [{"name": "Dish Name", "description": "brief desc", "price_cents": 1299}]
    }
  ]
}
RULES:
- price_cents = cents ($12.99 → 1299). Use 0 if price is not visible.
- Every category MUST have at least one item. No empty categories.
- Extract EVERY food item visible in the image.
- Include appetizers, mains, sides, drinks, desserts — everything visible.
- If no clear categories exist, group into logical categories (e.g. "Main Dishes", "Drinks", "Sides").
- For handwritten menus, do your best to read each item.
- Always return valid JSON."""


MIME_MAP = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".webp": "image/webp",
}


def _parse_json(raw_text: str) -> dict:
    """Parse JSON from AI response, handling ```json wrapping."""
    t = raw_text.strip()
    if t.startswith("```"):
        t = t.split("```", 2)[1]
        if t.startswith("json"):
            t = t[4:]
    return _json.loads(t)


def extract_menu_from_image(image_bytes: bytes, filename: str = "menu.jpg") -> dict:
    """
    Extract menu from an image using Gemini Vision API.

    Args:
        image_bytes: Raw image file bytes.
        filename: Original filename (for MIME type detection).

    Returns:
        Structured menu dict: {restaurant_name, categories: [{name, items}]}
    """
    load_dotenv()

    # Determine MIME type
    ext = os.path.splitext(filename)[1].lower()
    mime_type = MIME_MAP.get(ext, "image/jpeg")

    # Base64 encode
    image_b64 = base64.b64encode(image_bytes).decode("utf-8")

    menu_data = None

    # --- Try Gemini Vision first ---
    gemini_key = os.getenv("GEMINI_API_KEY", "")
    if gemini_key:
        try:
            gemini_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={gemini_key}"

            body = _json.dumps({
                "contents": [{
                    "parts": [
                        {"text": f"Look at this restaurant menu image carefully. {EXTRACTION_PROMPT}"},
                        {"inline_data": {"mime_type": mime_type, "data": image_b64}},
                    ]
                }],
                "generationConfig": {
                    "temperature": 0.1,
                    "responseMimeType": "application/json",
                }
            }).encode()

            req = urllib.request.Request(gemini_url, body, {"Content-Type": "application/json"})
            res = urllib.request.urlopen(req, timeout=120)
            data = _json.loads(res.read())
            text = data["candidates"][0]["content"]["parts"][0]["text"].strip()
            menu_data = _parse_json(text)
            print(f"[MenuExtractor] Gemini Vision extracted {sum(len(c.get('items', [])) for c in menu_data.get('categories', []))} items")
        except Exception as e:
            print(f"[MenuExtractor] Gemini Vision error: {e}")
            menu_data = None

    # --- Fallback to OpenAI GPT-4o-mini (vision) ---
    if not menu_data:
        oai_key = os.getenv("OPENAI_API_KEY", "")
        if oai_key:
            try:
                oai_url = "https://api.openai.com/v1/chat/completions"
                oai_body = _json.dumps({
                    "model": "gpt-4o-mini",
                    "response_format": {"type": "json_object"},
                    "messages": [
                        {"role": "system", "content": EXTRACTION_PROMPT},
                        {"role": "user", "content": [
                            {"type": "text", "text": "Extract the menu from this image."},
                            {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{image_b64}"}},
                        ]},
                    ],
                    "temperature": 0.1,
                    "max_tokens": 16000,
                }).encode()

                req = urllib.request.Request(oai_url, oai_body, {
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {oai_key}",
                })
                res = urllib.request.urlopen(req, timeout=120)
                data = _json.loads(res.read())
                text = data["choices"][0]["message"]["content"].strip()
                menu_data = _parse_json(text)
                print(f"[MenuExtractor] OpenAI extracted {sum(len(c.get('items', [])) for c in menu_data.get('categories', []))} items")
            except Exception as e:
                print(f"[MenuExtractor] OpenAI error: {e}")
                menu_data = None

    if not menu_data:
        raise ValueError("Failed to extract menu from image. Please try a clearer photo.")

    # Validate structure
    if "categories" not in menu_data:
        menu_data = {"categories": [], "restaurant_name": menu_data.get("restaurant_name", "Unknown")}

    # Ensure price_cents are integers
    for cat in menu_data.get("categories", []):
        for item in cat.get("items", []):
            try:
                item["price_cents"] = int(item.get("price_cents", 0))
            except (TypeError, ValueError):
                item["price_cents"] = 0
            if "description" not in item:
                item["description"] = ""

    return menu_data


# ---------------------------------------------------------------------------
# Document extraction: PDF, DOCX, XLSX → text → AI → structured menu
# ---------------------------------------------------------------------------

DOCUMENT_PROMPT = """You are extracting a restaurant menu from raw text extracted from a document (PDF, Word, or Excel).
The text may be messy, contain headers/footers, page numbers, or formatting artifacts.
Return JSON:
{
  "restaurant_name": "...",
  "categories": [
    {
      "name": "Category Name",
      "items": [{"name": "Dish Name", "description": "brief desc", "price_cents": 1299}]
    }
  ]
}
RULES:
- price_cents = cents ($12.99 → 1299). Use 0 if price is not visible.
- Every category MUST have at least one item. No empty categories.
- Extract EVERY food item found in the text.
- Include appetizers, mains, sides, drinks, desserts — everything.
- If no clear categories exist, group into logical categories.
- Ignore page numbers, headers, footers, website URLs, phone numbers.
- Always return valid JSON."""

DOCUMENT_MIME_MAP = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".xlsx": "xlsx",
    ".xls": "xlsx",
}


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pypdf, fallback to pdfplumber."""
    import tempfile
    pages_text = []

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        # Method 1: pypdf (fast)
        try:
            from pypdf import PdfReader
            reader = PdfReader(tmp_path, strict=False)
            for page in reader.pages:
                try:
                    text = page.extract_text()
                    if text and len(text.strip()) > 10:
                        pages_text.append(text.strip())
                except Exception:
                    pass
        except Exception as e:
            print(f"[MenuExtractor] pypdf failed: {e}")

        # Method 2: pdfplumber fallback
        if not pages_text:
            print("[MenuExtractor] pypdf yielded no text, trying pdfplumber...")
            try:
                import pdfplumber
                with pdfplumber.open(tmp_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text and text.strip():
                            pages_text.append(text.strip())
            except Exception as e:
                print(f"[MenuExtractor] pdfplumber failed: {e}")
    finally:
        os.unlink(tmp_path)

    return "\n\n".join(pages_text)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    import tempfile
    from docx import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n".join(paragraphs)
    finally:
        os.unlink(tmp_path)


def _extract_text_from_excel(file_bytes: bytes) -> str:
    """Extract text from Excel using openpyxl."""
    import tempfile
    import openpyxl

    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        wb = openpyxl.load_workbook(tmp_path, data_only=True)
        rows_text = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows_text.append(f"--- {sheet_name} ---")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell) for cell in row if cell is not None]
                if cells:
                    rows_text.append(" | ".join(cells))

        return "\n".join(rows_text)
    finally:
        os.unlink(tmp_path)


def extract_menu_from_document(file_bytes: bytes, filename: str = "menu.pdf") -> dict:
    """
    Extract menu from a document (PDF, DOCX, XLSX) using text extraction + AI.

    Pipeline: File → Text Extraction → Gemini AI → Structured JSON

    Args:
        file_bytes: Raw document file bytes.
        filename: Original filename (for type detection).

    Returns:
        Structured menu dict: {restaurant_name, categories: [{name, items}]}
    """
    load_dotenv()

    ext = os.path.splitext(filename)[1].lower()
    doc_type = DOCUMENT_MIME_MAP.get(ext)

    if not doc_type:
        raise ValueError(f"Unsupported document type: {ext}. Supported: PDF, DOCX, XLSX")

    # Step 1: Extract raw text
    print(f"[MenuExtractor] Extracting text from {doc_type}: {filename}")

    if doc_type == "pdf":
        raw_text = _extract_text_from_pdf(file_bytes)
    elif doc_type == "docx":
        raw_text = _extract_text_from_docx(file_bytes)
    elif doc_type == "xlsx":
        raw_text = _extract_text_from_excel(file_bytes)
    else:
        raise ValueError(f"Unsupported document type: {doc_type}")

    if not raw_text or len(raw_text.strip()) < 20:
        raise ValueError("Could not extract enough text from the document. The file may be empty or a scanned PDF (try uploading as an image instead).")

    print(f"[MenuExtractor] Extracted {len(raw_text)} chars of text")

    # Truncate to avoid token limits (keep first ~8000 chars which is plenty for menus)
    if len(raw_text) > 8000:
        raw_text = raw_text[:8000]
        print("[MenuExtractor] Text truncated to 8000 chars")

    # Step 2: Send to AI for menu parsing
    menu_data = None

    # --- Try Gemini first ---
    gemini_key = os.getenv("GEMINI_API_KEY", "")
    if gemini_key:
        try:
            gemini_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={gemini_key}"

            body = _json.dumps({
                "contents": [{
                    "parts": [
                        {"text": f"{DOCUMENT_PROMPT}\n\nHere is the raw menu text extracted from the document:\n\n{raw_text}"}
                    ]
                }],
                "generationConfig": {
                    "temperature": 0.1,
                    "responseMimeType": "application/json",
                }
            }).encode()

            req = urllib.request.Request(gemini_url, body, {"Content-Type": "application/json"})
            res = urllib.request.urlopen(req, timeout=120)
            data = _json.loads(res.read())
            text = data["candidates"][0]["content"]["parts"][0]["text"].strip()
            menu_data = _parse_json(text)
            print(f"[MenuExtractor] Gemini extracted {sum(len(c.get('items', [])) for c in menu_data.get('categories', []))} items from document")
        except Exception as e:
            print(f"[MenuExtractor] Gemini error: {e}")
            menu_data = None

    # --- Fallback to OpenAI ---
    if not menu_data:
        oai_key = os.getenv("OPENAI_API_KEY", "")
        if oai_key:
            try:
                oai_url = "https://api.openai.com/v1/chat/completions"
                oai_body = _json.dumps({
                    "model": "gpt-4o-mini",
                    "response_format": {"type": "json_object"},
                    "messages": [
                        {"role": "system", "content": DOCUMENT_PROMPT},
                        {"role": "user", "content": f"Extract the menu from this text:\n\n{raw_text}"},
                    ],
                    "temperature": 0.1,
                    "max_tokens": 16000,
                }).encode()

                req = urllib.request.Request(oai_url, oai_body, {
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {oai_key}",
                })
                res = urllib.request.urlopen(req, timeout=120)
                data = _json.loads(res.read())
                text = data["choices"][0]["message"]["content"].strip()
                menu_data = _parse_json(text)
                print(f"[MenuExtractor] OpenAI extracted {sum(len(c.get('items', [])) for c in menu_data.get('categories', []))} items from document")
            except Exception as e:
                print(f"[MenuExtractor] OpenAI error: {e}")
                menu_data = None

    if not menu_data:
        raise ValueError("Failed to extract menu from document. Please try a different file format.")

    # Validate structure (same as image extraction)
    if "categories" not in menu_data:
        menu_data = {"categories": [], "restaurant_name": menu_data.get("restaurant_name", "Unknown")}

    for cat in menu_data.get("categories", []):
        for item in cat.get("items", []):
            try:
                item["price_cents"] = int(item.get("price_cents", 0))
            except (TypeError, ValueError):
                item["price_cents"] = 0
            if "description" not in item:
                item["description"] = ""

    return menu_data
